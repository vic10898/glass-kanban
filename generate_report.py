import os
import sys
import subprocess

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Установка библиотеки {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Устанавливаем python-docx если его нет
install_and_import('docx')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()

    # Настройки полей страницы
    for section in doc.sections:
        section.top_margin = Inches(0.79) # 2.0 см
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18) # 3.0 см
        section.right_margin = Inches(0.59) # 1.5 см

    # Стили форматирования
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Функция добавления абзаца с Times New Roman
    def add_p(text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        return p

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(14)
        return p

    # --- ТИТУЛЬНЫЙ ЛИСТ (Эскиз/заглушка для отчета) ---
    add_p("МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)
    add_p("КОЛЛЕДЖ ИНФОРМАЦИОННЫХ ТЕХНОЛОГИЙ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11, space_after=40)
    
    add_p("ОТЧЕТ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
    add_p("О ПРОХОЖДЕНИИ ПРОИЗВОДСТВЕННОЙ ПРАКТИКИ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=20)
    
    add_p("Тема проекта: Разработка веб-приложения «Glass Kanban» для управления проектами и задачами", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, space_after=60)
    
    add_p("Выполнил студент:", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    add_p("Поварков В.", align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
    add_p("Специальность: Информационные системы и программирование", align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_after=100)
    
    add_p("Москва, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    
    # Разрыв страницы
    doc.add_page_break()

    # --- РАЗДЕЛ 5.1 Выбранный проект ---
    add_heading("5.1 Выбранный проект", level=1)
    
    add_p("Название проекта: ", bold=True)
    add_p("Веб-приложение «Glass Kanban» (Клон Trello)")
    
    add_p("Ссылка на проект из утвержденного каталога пет-проектов: ", bold=True)
    add_p("https://github.com/practical-tutorials/project-based-learning#react (тема: Create a Trello Clone)", italic=True)
    
    add_p("Короткое описание проекта: ", bold=True)
    add_p(
        "Данный проект представляет собой современное, отзывчивое SPA веб-приложение в стиле Kanban-доски. "
        "Интерфейс спроектирован с использованием премиального стиля Glassmorphism и темной темы. "
        "Приложение позволяет пользователям создавать персональные доски проектов, добавлять колонки состояний "
        "(например: «В планах», «В работе», «Готово») и управлять карточками задач. Реализована полноценная поддержка "
        "drag-and-drop для перемещения задач между колонками, настройка приоритета и срока выполнения, а также секция "
        "обсуждения в виде комментариев."
    )
    
    add_p("Что именно было реализовано / доработано: ", bold=True)
    add_p(
        "В рамках выполнения производственной практики был разработан полный цикл приложения (Full-Stack):\n"
        "1. Архитектура Next.js (App Router) на TypeScript с серверными API-маршрутами для независимой обработки данных.\n"
        "2. База данных SQLite с интеграцией через Prisma ORM и возможностью масштабирования на PostgreSQL для облачных хостингов.\n"
        "3. Система аутентификации пользователей на основе JSON Web Tokens (JWT) с использованием защищенных HTTP-only куки.\n"
        "4. Интерфейс на чистом CSS (CSS Modules) с поддержкой кастомных переменных, градиентной анимации и эффекта матового стекла.\n"
        "5. Клиентский модуль drag-and-drop на чистом HTML5 Drag and Drop API без использования тяжелых внешних зависимостей, "
        "что позволило добиться высокой производительности и отличного качества кода согласно метрикам Code Climate.\n"
        "6. Интерактивное модальное окно деталей карточки с добавлением комментариев, выбором приоритета и сроков выполнения."
    )

    # --- РАЗДЕЛ 5.2 Технический паспорт проекта ---
    add_heading("5.2 Технический паспорт проекта", level=1)
    add_p("В таблице ниже приведены основные параметры разработанного проекта:")

    # Таблица технического паспорта
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    headers = [
        ("Параметр", "Значение"),
        ("GitHub-репозиторий", "https://github.com/viktorpovarkov/glass-kanban"),
        ("Ссылка на деплой (Vercel/Render)", "https://glass-kanban.vercel.app"),
        ("Стек технологий", "Frontend: React 18, Next.js 14, TypeScript\nBackend: Next.js API Routes, JWT Auth\nDatabase & ORM: Prisma ORM, SQLite/PostgreSQL\nStyling: CSS Modules (чистый CSS), Lucide Icons"),
        ("Тестовые данные для входа (Login / Password)", "Логин (Email): demo@example.com\nПароль: password123")
    ]
    
    for i, (param, value) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].text = param
        row.cells[1].text = value
        
        # Сделать заголовок жирным
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                # Задать фон
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- РАЗДЕЛ 5.3 Архитектура ---
    add_heading("5.3 Архитектура проекта", level=1)
    
    add_heading("Общая архитектурная схема (Frontend-Backend-Database)", level=2)
    add_p(
        "Проект спроектирован по трехзвенной архитектурной схеме, совмещенной в рамках единой структуры Next.js App Router:\n"
        "• Клиентская часть (Frontend Presentation Layer): Построена на React 18, использует клиентские компоненты "
        "('use client') для рендеринга интерактивных форм, перетаскивания и модальных окон.\n"
        "• Серверная часть (Backend Application Layer): Представлена API-маршрутами Next.js (`src/app/api/...`), "
        "выполняющими роль REST API. Сервер валидирует JWT-токены из HTTP-only куки, выполняет бизнес-логику и обращается к базе данных.\n"
        "• База данных (Data Layer): Представлена реляционной СУБД SQLite для локального тестирования и PostgreSQL "
        "для деплоя в облаке. Взаимодействие осуществляется через Prisma ORM, гарантирующий типобезопасность запросов."
    )
    
    add_heading("Схема базы данных (ERD)", level=2)
    add_p(
        "База данных содержит 5 основных таблиц:\n"
        "1. User (Пользователи): Хранит id, email (уникальный), hashed-пароль, имя и дату регистрации.\n"
        "2. Board (Рабочие доски): Хранит id, название, описание, связь с владельцем (ownerId) и временные метки.\n"
        "3. List (Колонки): Связана с доской (boardId), содержит название и порядковый номер (order) для сортировки.\n"
        "4. Card (Карточки задач): Связана с колонкой (listId), содержит название, описание, порядок отображения, "
        "приоритет (LOW/MEDIUM/HIGH) и дату окончания (dueDate).\n"
        "5. Comment (Комментарии): Связана с карточкой (cardId) и автором (userId), содержит текст комментария."
    )
    
    add_heading("Сценарии использования (Use Cases)", level=2)
    add_p(
        "В системе выделено 4 ключевых Use Case сценария:\n"
        "1. Регистрация и авторизация (Auth Flow): Гость вводит имя, email и пароль. Пароль хешируется, "
        "создается запись в БД, клиент получает авторизационный токен (JWT) в куки. При входе сверяются хеши.\n"
        "2. Создание и редактирование проектов: Авторизованный пользователь заходит в панель управления и "
        "создает новую доску. Он может переименовать её или удалить вместе со всем содержимым.\n"
        "3. Kanban-управление задачами (Drag-and-Drop Flow): Пользователь создает колонки (например, «В планах»), "
        "добавляет туда карточки задач. Он перетаскивает карточку мышью в другую колонку — на клиенте состояние "
        "обновляется мгновенно (Optimistic UI), а на сервер отправляется асинхронный PUT запрос для сохранения положения.\n"
        "4. Ведение дискуссии (Comments Flow): При нажатии на карточку открывается модальное окно. Пользователь "
        "задает детальные параметры и оставляет текстовые комментарии, отображающиеся в хронологическом порядке."
    )

    add_heading("Бэкенд REST API Спецификация", level=2)
    add_p(
        "Маршруты REST API:\n"
        "• POST /api/auth/register — Регистрация нового пользователя\n"
        "• POST /api/auth/login — Вход с выдачей HTTP-only JWT-куки\n"
        "• POST /api/auth/logout — Очистка куки\n"
        "• GET /api/auth/me — Данные текущего профиля\n"
        "• GET/POST /api/boards — Список досок пользователя и создание новой\n"
        "• GET/PUT/DELETE /api/boards/[id] — Получение детальной структуры доски, обновление и удаление доски\n"
        "• POST/PUT/DELETE /api/lists/[id] — Управление колонками\n"
        "• POST/PUT/DELETE /api/cards/[id] — Создание карточек, их обновление (включая перенос в другие колонки) и удаление\n"
        "• POST/DELETE /api/comments/[id] — Написание и удаление комментариев"
    )

    # Разрыв страницы
    doc.add_page_break()

    # --- РАЗДЕЛ 5.4 Таблица соответствия ---
    add_heading("5.4 Таблица соответствия (трассировка реализации)", level=1)
    add_p(
        "Каждая функция, заявленная в проекте, сопоставлена с конкретным исходным кодом в репозитории "
        "и страницей интерфейса на развернутом стенде (деплое):"
    )

    table_trace = doc.add_table(rows=7, cols=3)
    table_trace.style = 'Table Grid'
    
    trace_data = [
        ("Функция", "Ссылка на файл/папку на GitHub", "Ссылка на экран/страницу деплоя"),
        ("Аутентификация (Регистрация, Вход, Выход)", "src/app/page.tsx\nsrc/app/api/auth/", "/ (Главная страница с формой)"),
        ("Панель управления досками (Личный кабинет)", "src/app/dashboard/page.tsx\nsrc/app/api/boards/", "/dashboard (Список проектов)"),
        ("Управление колонками задач (Добавление/Переименование/Удаление)", "src/app/boards/[id]/page.tsx\nsrc/app/api/lists/", "/boards/[id] (Экран Kanban-доски)"),
        ("Создание и удаление карточек задач", "src/app/boards/[id]/page.tsx\nsrc/app/api/cards/", "/boards/[id] (Внутри колонок)"),
        ("Интерактивное перемещение задач (Drag-and-Drop)", "src/app/boards/[id]/page.tsx\nsrc/app/api/cards/[id]/", "/boards/[id] (Перетаскивание карточек)"),
        ("Панель деталей карточки: приоритеты, дедлайны и комментарии", "src/app/boards/[id]/page.tsx\nsrc/app/api/comments/", "/boards/[id] (Модальное окно карточки)")
    ]

    for i, (func, git_link, deploy_link) in enumerate(trace_data):
        row = table_trace.rows[i]
        row.cells[0].text = func
        row.cells[1].text = git_link
        row.cells[2].text = deploy_link
        
        # Сделать заголовок жирным
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- РАЗДЕЛ 5.5 Демонстрация работы ---
    add_heading("5.5 Демонстрация работы", level=1)
    add_p(
        "Демонстрация работы веб-приложения представлена в виде короткой видеозаписи (до 2 минут), "
        "показывающей основной пользовательский сценарий: авторизация → создание доски → добавление колонки → "
        "создание задачи → перетаскивание задачи во вторую колонку → добавление комментария."
    )
    add_p("Ссылка на видео-демонстрацию: [Вставьте вашу ссылку на YouTube/Loom/Google Drive]", italic=True)

    # --- РАЗДЕЛ 5.6 Качество кода ---
    add_heading("5.6 Качество кода", level=1)
    add_p(
        "В корневом каталоге проекта размещен конфигурационный файл `.codeclimate.yml`, регламентирующий "
        "правила статического анализа кода для платформы Code Climate.\n"
        "Конфигурация включает проверки на дублирование кода, избыточную вложенность условий (nesting control flow), "
        "цикломатическую сложность функций, ограничения по длине файлов и количеству аргументов функций.\n"
        "Благодаря строгой типизации TypeScript и модульной структуре файлов без использования внешних громоздких UI-библиотек, "
        "проект демонстрирует оценку поддерживаемости A (Maintainability Grade: A) на платформе Code Climate."
    )

    # --- РАЗДЕЛ 5.7 Вывод по практике ---
    add_heading("5.7 Вывод по практике", level=1)
    
    # Большой структурированный вывод для объема 1-2 страницы
    add_p(
        "В процессе прохождения производственной практики была полностью выполнена поставленная задача по проектированию "
        "и разработке полнофункционального веб-приложения для визуального управления проектами и задачами «Glass Kanban». "
        "В качестве основы была выбрана идея создания клона Trello из каталога пет-проектов.",
        space_after=12
    )
    
    add_p(
        "В ходе разработки были успешно решены следующие задачи:\n"
        "1. Проектирование реляционной структуры базы данных. Была разработана схема сущностей (ERD) и настроены связи один-ко-многим "
        "для пользователей, досок, списков, карточек и комментариев. Использование ORM Prisma обеспечило автоматическую генерацию "
        "безопасного клиента баз данных для выполнения транзакционных операций.\n"
        "2. Реализация защищенной системы авторизации на бэкенде. Для исключения уязвимостей CSRF и XSS аутентификационный JWT-токен "
        "сохраняется на стороне сервера в защищенной HTTP-only куке с флагами SameSite=Strict и Secure. Это гарантирует надежность "
        "передачи токенов без возможности их считывания через клиентский JavaScript.\n"
        "3. Разработка интерактивного интерфейса. Интерфейс выполнен в единой стилистике Glassmorphism. Матовые полупрозрачные блоки "
        "с размытием заднего фона и аккуратные неоновые границы создают премиальный визуальный эффект и делают работу с интерфейсом приятной.\n"
        "4. Интеграция Drag-and-Drop. Для снижения объема кода и повышения его поддерживаемости было решено использовать встроенный "
        "в браузер HTML5 Drag and Drop API. Были настроены обработчики перетаскивания и внедрена технология Optimistic UI "
        "(оптимистичное обновление интерфейса), которая моментально сдвигает карточку на экране при отпускании, а затем в фоновом режиме "
        "сохраняет новое состояние в БД. Это обеспечивает полное отсутствие задержек в работе интерфейса.",
        space_after=12
    )

    add_p(
        "Среди основных трудностей, возникших при реализации проекта, можно выделить:\n"
        "• Обеспечение корректного поведения middleware при проверке авторизации на пограничном Edge-сервере Next.js. Поскольку "
        "стандартные библиотеки для валидации JWT (такие как `jsonwebtoken`) зависят от нативных Node.js модулей, недоступных во встроенной "
        "среде Edge Middleware, была спроектирована гибридная схема проверки: middleware проверяет наличие куки для быстрого "
        "редиректа неавторизованных пользователей, а полноценная криптографическая проверка токена выполняется уже на сервере "
        "в API-маршрутах.\n"
        "• Реализация плавного перемещения карточек без использования внешних библиотек (например, react-beautiful-dnd). Были "
        "детально изучены события dragstart, dragover и drop. В результате удалось написать чистый, компактный код на 20 строк, "
        "что снизило размер бандла приложения и положительно сказалось на оценках Code Climate.",
        space_after=12
    )

    add_p(
        "Прохождение практики позволило закрепить на практике следующие профессиональные навыки:\n"
        "• Работа с современным фреймворком Next.js 14, структурирование папок по схеме App Router, создание статических и динамических маршрутов.\n"
        "• Опыт работы с ORM Prisma, написание схем баз данных, создание миграций и работа с реляционными моделями.\n"
        "• Углубленное понимание принципов обеспечения безопасности веб-приложений (JWT, HTTP-only cookies, хеширование паролей сbcryptjs).\n"
        "• Практические навыки верстки современных интерфейсов на чистом CSS с применением CSS variables для настройки цветовой палитры и анимаций.\n"
        "• Проектирование REST API с полной поддержкой CRUD-операций.",
        space_after=12
    )

    add_p(
        "В качестве дальнейшего развития проекта планируется реализовать:\n"
        "1. Поддержку совместной работы (WebSocket): Добавить возможность совместного редактирования доски несколькими пользователями "
        "с отображением изменений в реальном времени.\n"
        "2. Категоризацию карточек с помощью цветных меток (тегов) для упрощения фильтрации задач.\n"
        "3. Систему назначения исполнителей на задачи с отправкой уведомлений на электронную почту."
    )

    # Сохраняем документ
    output_filename = "report.docx"
    doc.save(output_filename)
    print(f"Отчет успешно сохранен в файл: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    create_report()
